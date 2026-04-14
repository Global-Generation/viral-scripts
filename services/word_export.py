"""Generate Word (.docx) export of a script."""
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_script_docx(script, video=None) -> io.BytesIO:
    """Generate a Word document for a script and return as BytesIO."""
    doc = Document()

    # Title
    title = doc.add_heading(level=1)
    run = title.add_run(f"Script #{script.id}")
    run.font.size = Pt(20)
    if video and video.title:
        sub = doc.add_paragraph()
        sub_run = sub.add_run(video.title)
        sub_run.font.size = Pt(12)
        sub_run.font.color.rgb = RGBColor(107, 114, 128)

    # Metadata table
    meta_table = doc.add_table(rows=0, cols=2)
    meta_table.style = "Light Grid"
    meta_items = [
        ("Channel", (script.channel or "—").upper()),
        ("Assigned To", (script.assigned_to or "—").capitalize()),
        ("Character", (script.character_type or "—").capitalize()),
        ("Viral Score", str(script.viral_score or "—")),
        ("Status", script.status or "—"),
    ]
    for label, value in meta_items:
        row = meta_table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value

    doc.add_paragraph()  # spacer

    # Final text (if available)
    if script.final_text:
        doc.add_heading("Final Script", level=2)
        doc.add_paragraph(script.final_text)
        doc.add_paragraph()

    # Pipeline stages (if available)
    if hasattr(script, "pipeline_stages") and script.pipeline_stages:
        accepted = [s for s in script.pipeline_stages if s.status == "accepted"]
        if accepted:
            stage_labels = {"intro": "Intro", "part1": "Part 1", "part2": "Part 2", "part3": "Part 3"}
            for stage in accepted:
                label = stage_labels.get(stage.stage_name, stage.stage_name)
                doc.add_heading(label, level=2)
                doc.add_paragraph(stage.result_text or "")
            doc.add_paragraph()

    # Original script
    if script.original_text:
        doc.add_heading("Original Script", level=2)
        doc.add_paragraph(script.original_text)

    # Provocative version
    if script.modified_text:
        doc.add_heading("Provocative Version", level=2)
        doc.add_paragraph(script.modified_text)

    # Video prompts
    for i, prompt_field in enumerate(["video1_prompt", "video2_prompt"], 1):
        prompt = getattr(script, prompt_field, "")
        if prompt:
            doc.add_heading(f"Video {i} Prompt", level=2)
            doc.add_paragraph(prompt)

    # Fact-check summary
    if script.fact_check_report:
        doc.add_heading("Fact-Check Report", level=2)
        import json
        try:
            report = json.loads(script.fact_check_report)
            for fact in report.get("facts", []):
                p = doc.add_paragraph()
                status_map = {"verified": "VERIFIED", "false": "FALSE",
                              "partial": "PARTIALLY VERIFIED", "unverifiable": "UNVERIFIABLE"}
                status = status_map.get(fact.get("verdict", ""), fact.get("verdict", ""))
                p.add_run(f"[{status}] ").bold = True
                p.add_run(fact.get("claim", ""))
        except (json.JSONDecodeError, TypeError):
            doc.add_paragraph(script.fact_check_report)

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
